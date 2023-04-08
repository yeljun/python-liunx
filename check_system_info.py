#!./venv/bin/python3
"""
author:ranyijun@uniontech.com
date:2023-3-11
scription:check system info
ver:v1.1:
"""


from docx import Document
import os

def get_res_comd(command):
    res = os.popen(command).readlines()
    return res



def get_host_ip():
    val = os.popen("ip a s | awk '/inet /{print $2}' | grep -v 127.0.0.1 |awk -F'/' '{print $1}'").readlines()
    #print(val)
    name_ip = val[0]
    name_ip = name_ip[0:-1]

    print(name_ip)
    return name_ip



def set_docx_text(document,scription,command2):
    command1 = scription
    command2 = command2
    get_res_res = get_res_comd(command2)
    list1=[]
    for i in range(len(get_res_res)):
        text=get_res_res[i]
        text=text[0:-1]
        list1.append(text)
    get_res_res = list1
    for m in range(len(get_res_res)):
        if len(get_res_res)==1:
            p = document.add_paragraph('%s%s' % (command1 + ":", get_res_res[m]))
        else:
            if m == 0:
                p = document.add_paragraph('%s%s' % (command1+":\n", get_res_res[m]))
            if m != 0:
                p = document.add_paragraph('%s' % (get_res_res[m]))


def set_docx_acl_text(document,scription,command2,acl):
    command1 = scription
    command2 = command2
    get_res_res = get_res_comd(command2)
    acl_set_var = acl
    if get_res_res == acl_set_var:
        #print("zhengchanbg")
        p = document.add_paragraph('%s%s' % (command1 + ":", "正常"))
    else:
        p = document.add_paragraph('%s%s' % (command1 + ":", "异常"))

        list1=[]
        for i in range(len(get_res_res)):
            text=get_res_res[i]
            text=text[0:-1]
            list1.append(text)
        get_res_res = list1
        for m in range(len(get_res_res)):
            if len(get_res_res)==1:
                p = document.add_paragraph('%s%s' % (command1 + ":", get_res_res[m]))
            else:
                if m == 0:
                    p = document.add_paragraph('%s%s' % (command1+":\n", get_res_res[m]))
                if m != 0:
                    p = document.add_paragraph('%s' % (get_res_res[m]))



def new_docx():
    document = Document()


    val=get_host_ip()
    #添加标题，并设置级别，范围：0 至 9，默认为1
    print(val)
    document.add_heading('%s服务器巡检报告' %val, level=2)


    '''
    获取os信息
    '''
    #p = document.add_paragraph('获取os信息',level=3)
    document.add_paragraph('获取os信息', style='Intense Quote')
    #添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
    set_docx_text(document,"内核信息","uname -r")
    set_docx_text(document, "操作系统版本", "cat /etc/os-version")
    set_docx_text(document, "检查时间", "uptime | awk -F \",\" '{print $1}'")
    set_docx_text(document, "用户登陆数", "uptime | awk -F \",\" '{print $3}'")


    '''
    获取内存相关信息
    '''
    document.add_paragraph('获取内存相关信息', style='Intense Quote')
    set_docx_text(document, "总内存容量", "free -m | awk '/Mem/{printf \"%.2fG\", $2/1024}'")
    set_docx_text(document, "用户程序占用内存量", "free -m | awk '/Mem/{printf \"%.2fG\", $3/1024}'")
    set_docx_text(document, "多进程共享占用内存量", "free -m | awk '/Mem/{printf \"%.2fG\", $5/1024}'")
    set_docx_text(document, "缓存占用内存量", "free -m | awk '/Mem/{printf \"%.2fG\", $2/1024}'")
    set_docx_text(document, "空闲内存容量", "free -m | awk '/Mem/{printf \"%.2fG\", $4/1024 }'")
    set_docx_text(document, "剩余可用内存容量", "free -m | awk 'NR==2{printf \"%.2fG\",$NF/1024}'")
    set_docx_text(document, "可用内存使用占比", "free -m | awk '/Mem/{printf \"%.2f\", $NF/$2*100}'")
    set_docx_text(document, "总的交换分区容量", "free -m | awk '/Swap/{printf \"%.2fG\", $2/1024}'")
    set_docx_text(document, "用户使用的交换分区容量", "free -m | awk '/Swap/{printf \"%.2fG\",$3/1024}'")
    set_docx_text(document, "剩余交换分区容量", "free -m | awk '/Swap/{printf \"%.2fG\", $4/1024}'")
    set_docx_text(document, "可用交换分区占比", "free -m | awk '/Swap/{printf \"%.2f\",$4/$2*100}'")
    set_docx_text(document, "总占用内存资源最多的10个进程列表", "ps --no-headers -eo rss,args | sort -k1 -n -r | head -10")

    '''
    获取cpu相关信息
    '''
    document.add_paragraph('获取cpu相关信息', style='Intense Quote')
    set_docx_text(document, "CPU型号", "LANG=C lscpu | awk -F: '/Model name/ {print $2}'")
    set_docx_text(document, "CPU内核数量", "awk '/processor/{core++} END{print core}' /proc/cpuinfo")
    set_docx_text(document, "CPU最近1/5/15分钟的平均负载", "uptime | sed 's/,/ /g' | awk '{for(i=NF-2;i<=NF;i++)print $i }' | xargs")
    set_docx_text(document, "发生中断数量", "vmstat 1 1 | awk 'NR==3{print $11}'")
    set_docx_text(document, "上下文切换数量", "vmstat 1 1 | awk 'NR==3{print $12}'")
    set_docx_text(document, "占用CPU资源最多的10个进程列表", "ps --no-headers -eo %cpu,args | sort -k1 -n -r | head -10")

    '''
    获取disk_io相关信���
    '''
    document.add_paragraph('获取disk_io相关信息', style='Intense Quote')
    set_docx_text(document, "分区挂载信息", "df -h")
    set_docx_text(document, "磁盘总数", "vmstat -D | awk 'NR==1{print $1}'")
    set_docx_text(document, "分区总数", "vmstat -D | awk 'NR==2{print $1}'")
    set_docx_text(document, "磁盘分区信息", "lsblk -n")
    set_docx_text(document, "写入磁盘的总块数", "vmstat 1 1 | awk 'NR==3{print $10}'")
    set_docx_text(document, "从磁盘读出的块数", "vmstat 1 1 | awk 'NR==3{print $9}'")
    set_docx_text(document, "每秒写磁盘块的速率", "vmstat 1 1 | awk 'NR==3{print $16}'")

    '''
    获取process相关信息
    '''

    document.add_paragraph('获取process相关信息', style='Intense Quote')
    set_docx_text(document, "进程数量", "ps aux | wc -l")
    set_docx_text(document, "用户的最大进程数", "bash -c ulimit -u")
    set_docx_text(document, "内核设置的最大进程数", "sysctl kernel.pid_max | awk '{print $3}'")
    set_docx_text(document,"僵尸进程数量","ps -A -ostat,ppid,pid,cmd | grep -e '^[Zz]'|wc -l")

    '''
    获取network相关信息
    '''
    document.add_paragraph('获取network相关信息', style='Intense Quote')
    set_docx_text(document, "获取网卡流量信息", "cat /proc/net/dev | tail -n +3 | awk 'BEGIN{ print \"网卡名称 入站数据流量(bytes) 出站数据流量(bytes)\"} {print $1,$2,$10}' | column -t")
    set_docx_text(document, "获取暴露端口信息", "ss -ntulpa")
    set_docx_text(document, "本地IP地址列表", "ip a s | awk '/inet /{print $2}'")

    '''
    获取系统重要目录权限相关信息
    '''
    document.add_paragraph('获取系统重要目录权限相关信息', style='Intense Quote')
    acl_set_var=['# file: var\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::r-x\n', 'other::r-x\n', '\n']
    set_docx_acl_text(document, "/var目录权限", "getfacl /var ", acl_set_var)

    acl_set_root = ['# file: root\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::---\n', 'mask::--x\n', 'other::---\n',
                   '\n']
    set_docx_acl_text(document, "/root目录权限", "getfacl /root ", acl_set_root)

    acl_set_home = ['# file: home\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::---\n', 'mask::--x\n', 'other::---\n',
                   '\n']
    set_docx_acl_text(document, "/home目录权限", "getfacl /home ", acl_set_home)

    acl_set_etc = ['# file: etc\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::r-x\n', 'other::r-x\n',
                   '\n']
    set_docx_acl_text(document, "/etc目录权限", "getfacl /etc ", acl_set_etc)

    acl_set_data = ['# file: data\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::r-x\n', 'other::r-x\n',
                   '\n']
    set_docx_acl_text(document, "/data目录权限", "getfacl /data ", acl_set_data)

    acl_set_bin = ['# file: bin\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::r-x\n', 'other::r-x\n',
                   '\n']
    set_docx_acl_text(document, "/bin目录权限", "getfacl /bin ", acl_set_bin)

    acl_set_usr = ['# file: usr\n', '# owner: root\n', '# group: root\n', 'user::rwx\n', 'group::r-x\n', 'other::r-x\n',
                   '\n']
    set_docx_acl_text(document, "/usr目录权限", "getfacl /usr ", acl_set_usr)


    document.add_page_break()
    document.save('check_system_info.docx')

if __name__=='__main__':
    new_docx()








